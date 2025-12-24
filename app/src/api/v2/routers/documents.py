from typing import List
from uuid import UUID

import chardet
import pdfplumber
from dishka.integrations.fastapi import FromDishka, inject
from docx import Document as Document_docx
from fastapi import APIRouter, File, HTTPException, UploadFile, status, Form
from result import Result

from src.api.v2.exceptions import handle_exception
from src.api.v2.helpers import check_admin_rights, get_user_id
from src.api.v2.schemas import Document, DocumentShort, DocumentUpdate, UploadDocumentRequest
from src.core.application.embeddings.schemas.load_data import LoadDataRequest
from src.core.application.embeddings.use_cases.load_data import LoadingUC
from src.core.application.rdb.schemas import DocumentResponse, UpdateDocumentRequest
from src.core.application.rdb.use_cases import (
    CheckAdminUC,
    DeleteDocumentUC,
    GetAvailableDocumentsUC,
    GetDocumentUC,
    UpdateDocumentUC,
)
from src.core.application.rdb.use_cases.user import RegisterUserUC
from src.core.domain.document import CoreDocument

router = APIRouter(
    prefix="/collections/{collection_id}/documents", tags=["documents"])

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB


@router.get("/", response_model=List[DocumentShort])
@inject
async def get_documents(
    telegram_id: int,
    collection_id: UUID,
    register_user_uc: FromDishka[RegisterUserUC],
    get_avaliable_documents_uc: FromDishka[GetAvailableDocumentsUC]
):
    """
    Получить список доступных документов.
    """
    try:
        user_id = await get_user_id(telegram_id, None, register_user_uc)
        documents: List[DocumentResponse] = await get_avaliable_documents_uc.execute(collection_id, user_id)

        return [
            DocumentShort(
                document_id=doc.document_id,
                collection_id=doc.collection_id,
                name=doc.title
            )
            for doc in documents
        ]
    except Exception as exc:
        raise handle_exception(exc) from None


@router.post("/", response_model=DocumentShort,
             status_code=status.HTTP_201_CREATED)
@inject
async def upload_document(
    register_user_uc: FromDishka[RegisterUserUC],
    check_admin_uc: FromDishka[CheckAdminUC],
    loading_uc: FromDishka[LoadingUC],
    telegram_id: str = Form(...),
    collection_id_new: str = Form(...),
    title: str = Form(...),
    description: str | None = Form(None),
    url: str | None = Form(None),
    file: UploadFile = File(...)
):
    """
    Загрузить новый документ.
    """
    try:
        user_id = await get_user_id(int(telegram_id), None, register_user_uc)
        await check_admin_rights(user_id, check_admin_uc)

        collection_id_new = str(collection_id_new)

        allowed_extensions = ['.txt', '.pdf', '.doc', '.docx']
        file_extension = file.filename.split(
            '.')[-1].lower() if '.' in file.filename else ''
        if f'.{file_extension}' not in allowed_extensions:
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail="Данный формат файлов не поддерживается"
            )

        file.file.seek(0, 2)
        file_size = file.file.tell()
        file.file.seek(0)

        if file_size > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"Размер файла превышает допустимый лимит в {MAX_FILE_SIZE / (1024*1024):.1f} МБ"
            )

        if f'.{file_extension}' == '.txt':
            content = await file.read()
            detected = chardet.detect(content)
            encoding = detected['encoding'] or 'utf-8'
            try:
                text = content.decode(encoding)
            except UnicodeDecodeError:
                text = content.decode('utf-8', errors='replace')
        elif f'.{file_extension}' == '.pdf':
            content = await file.read()
            with pdfplumber.open(content.decode()) as pdf:
                pages = [page.extract_text() for page in pdf.pages]
                text = '\n'.join(pages)
        elif f'.{file_extension}' in ['.doc', '.docx']:
            content = await file.read()
            if f'.{file_extension}' == '.docx':
                doc = Document_docx(content.decode())
                text = '\n'.join([para.text for para in doc.paragraphs])
            else:
                raise HTTPException(
                    status_code=status.HTTP_501_NOT_IMPLEMENTED,
                    detail="На данный момент файлы .doc не поддерживаются."
                )
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Bad request"
            )

        document = CoreDocument(
            text=text,
            metadata={
                "title": title,
                "description": description,
                "url": url,
                "filename": file.filename,
                "filesize": int(round(file_size / (1024 * 1024), 1))
            }
        )
        loading_request = LoadDataRequest(data=document)

        result: Result[str, str] = await loading_uc.execute(
            request=loading_request,
            collection_id=collection_id_new,
            title=title,
            file_name=file.filename,
            file_size=file_size
        )

        if result.is_err():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Ошибка загрузки документа: {result.err()}"
            )

        document_id = UUID(result.value)

        return DocumentShort(
            document_id=document_id,
            collection_id=collection_id_new,
            name=title
        )
    except Exception as exc:
        raise handle_exception(exc) from None


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
@inject
async def delete_document(
    telegram_id: int,
    document_id: UUID,
    register_user_uc: FromDishka[RegisterUserUC],
    check_admin_uc: FromDishka[CheckAdminUC],
    delete_document_uc: FromDishka[DeleteDocumentUC]
):
    """
    Удалить существующий документ.
    """
    try:
        user_id = await get_user_id(telegram_id, None, register_user_uc)
        await check_admin_rights(user_id, check_admin_uc)
        await delete_document_uc.execute(document_id, user_id)
    except Exception as exc:
        raise handle_exception(exc) from None


@router.get("/{document_id}", response_model=Document)
@inject
async def get_document(
    telegram_id: int,
    document_id: int,
    register_user_uc: FromDishka[RegisterUserUC],
    get_document_uc: FromDishka[GetDocumentUC]
):
    """
    Получить конкретный документ.
    """
    try:
        user_id = await get_user_id(telegram_id, None, register_user_uc)
        document: DocumentResponse = await get_document_uc.execute(document_id, user_id)

        return Document(
            document_id=document.document_id,
            collection_id=document.collection_id,
            title=document.title,
            file_name=document.file_name,
            file_size=document.file_size,
            created_at=document.created_at,
            metadata=document.metadata
        )
    except Exception as exc:
        raise handle_exception(exc) from None


@router.patch("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
@inject
async def update_document(
    document_data: DocumentUpdate,
    register_user_uc: FromDishka[RegisterUserUC],
    check_admin_uc: FromDishka[CheckAdminUC],
    update_document_uc: FromDishka[UpdateDocumentUC]
):
    """
    Редактировать конкретный документ.
    """
    try:
        user_id = await get_user_id(None, document_data, register_user_uc)
        await check_admin_rights(user_id, check_admin_uc)

        update_document_request = UpdateDocumentRequest(
            document_id=document_data.document_id,
            title=document_data.title
        )

        await update_document_uc.execute(update_document_request, user_id)
    except Exception as exc:
        raise handle_exception(exc) from None
